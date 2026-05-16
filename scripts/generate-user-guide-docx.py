from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
GUIDES_DIR = ROOT / "docs" / "user-guides"
TOP_LEVEL_OUTPUT_DIR = DOCS_DIR / "docx"
USER_GUIDE_OUTPUT_DIR = GUIDES_DIR / "docx"
POWERED_BY = "Powered by Viewertech"
POWERED_BY_URL = "https://viewertech.net"
TOP_LEVEL_DOCS = (
    "DEPLOYMENT.md",
    "CONFIGURATION.md",
    "HANDOVER.md",
    "SECURITY.md",
    "INTEGRATION_AUDIT_REPORT.md",
    "DUPLICATION_RISK_REPORT.md",
    "REFACTORING_PLAN.md",
    "CODE_IMPROVEMENTS.md",
    "FINAL_IMPLEMENTATION_SUMMARY.md",
)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "144870")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_cell_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def configure_document(document: Document, title: str):
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 24, RGBColor(20, 72, 112)),
        ("Heading 1", 20, RGBColor(20, 72, 112)),
        ("Heading 2", 15, RGBColor(15, 42, 60)),
        ("Heading 3", 12, RGBColor(120, 91, 27)),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display" if "Heading" in style_name or style_name == "Title" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    code_style = styles.add_style("EMS Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(23, 33, 27)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.18)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)

    document.core_properties.title = title
    document.core_properties.author = "Viewertech"
    document.core_properties.subject = "Embassy Management ERPNext user guide"


def add_runs_with_inline_markup(paragraph, text):
    link_pattern = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
    cursor = 0
    for match in link_pattern.finditer(text):
        add_formatted_runs(paragraph, text[cursor : match.start()])
        add_hyperlink(paragraph, match.group(1), match.group(2))
        cursor = match.end()
    add_formatted_runs(paragraph, text[cursor:])


def add_formatted_runs(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(20, 72, 112)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_footer(document: Document):
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Embassy Management ERPNext | ")
    add_hyperlink(paragraph, POWERED_BY, POWERED_BY_URL)


def markdown_to_docx(source: Path, destination: Path):
    lines = source.read_text(encoding="utf-8").splitlines()
    title = next((line[2:].strip() for line in lines if line.startswith("# ")), source.stem.replace("-", " ").title())

    document = Document()
    configure_document(document, title)
    add_footer(document)

    in_code = False
    code_language = ""

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.startswith("```"):
            in_code = not in_code
            code_language = line.strip("`").strip()
            if in_code and code_language:
                p = document.add_paragraph(style="EMS Code Block")
                set_cell_shading(p, "EEF4F8")
                p.add_run(code_language)
            continue

        if in_code:
            p = document.add_paragraph(style="EMS Code Block")
            set_cell_shading(p, "EEF4F8")
            p.add_run(line or " ")
            continue

        if not line:
            continue

        if line.startswith("# "):
            p = document.add_paragraph(style="Title")
            add_runs_with_inline_markup(p, line[2:].strip())
            continue

        if line.startswith("## "):
            p = document.add_paragraph(style="Heading 2")
            add_runs_with_inline_markup(p, line[3:].strip())
            continue

        if line.startswith("### "):
            p = document.add_paragraph(style="Heading 3")
            add_runs_with_inline_markup(p, line[4:].strip())
            continue

        bullet_match = re.match(r"^\s*-\s+(.*)", line)
        if bullet_match:
            p = document.add_paragraph(style="List Bullet")
            add_runs_with_inline_markup(p, bullet_match.group(1).strip())
            continue

        number_match = re.match(r"^\s*\d+\.\s+(.*)", line)
        if number_match:
            p = document.add_paragraph(style="List Number")
            add_runs_with_inline_markup(p, number_match.group(1).strip())
            continue

        p = document.add_paragraph()
        add_runs_with_inline_markup(p, line)

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def output_name_for(source: Path) -> str:
    if source.name.lower() == "readme.md":
        return "embassy-management-user-guides-index.docx"
    return f"{source.stem.lower()}.docx"


def main():
    batches = [
        (
            TOP_LEVEL_OUTPUT_DIR,
            [DOCS_DIR / filename for filename in TOP_LEVEL_DOCS],
        ),
        (
            USER_GUIDE_OUTPUT_DIR,
            sorted(GUIDES_DIR.glob("*.md"), key=lambda p: (p.name.lower() != "readme.md", p.name.lower()))
            if GUIDES_DIR.exists()
            else [],
        ),
    ]

    for output_dir, sources in batches:
        for source in sources:
            if not source.exists():
                continue
            destination = output_dir / output_name_for(source)
            markdown_to_docx(source, destination)
            print(destination.relative_to(ROOT))


if __name__ == "__main__":
    main()
